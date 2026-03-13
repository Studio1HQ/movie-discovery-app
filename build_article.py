"""
Build the updated article Word document.
- Unchanged text  : plain black
- Changed/new text: yellow highlight  (WD_COLOR_INDEX.YELLOW)
- Removed content : red strikethrough (noted inline)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.dml import MSO_THEME_COLOR
import copy

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def highlight_run(run, color="yellow"):
    """Add yellow highlight to a run via XML."""
    rPr = run._r.get_or_add_rPr()
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), color)
    rPr.append(hl)

def strike_run(run):
    """Add strikethrough to a run."""
    rPr = run._r.get_or_add_rPr()
    strike = OxmlElement("w:strike")
    strike.set(qn("w:val"), "true")
    rPr.append(strike)

def add_heading(text, level=1, changed=False):
    p = doc.add_heading(text, level=level)
    if changed:
        for run in p.runs:
            highlight_run(run)
    return p

def add_para(segments, style_name="Normal", align=None):
    """
    segments: list of (text, changed, bold, italic, strike, color)
    changed=True  → yellow highlight
    strike=True   → red strikethrough
    """
    p = doc.add_paragraph(style=style_name)
    if align:
        p.alignment = align
    for seg in segments:
        text   = seg[0]
        changed= seg[1] if len(seg) > 1 else False
        bold   = seg[2] if len(seg) > 2 else False
        italic = seg[3] if len(seg) > 3 else False
        strike = seg[4] if len(seg) > 4 else False
        color  = seg[5] if len(seg) > 5 else None
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        if changed:
            highlight_run(run)
        if strike:
            strike_run(run)
            run.font.color.rgb = RGBColor(192, 0, 0)
    return p

def add_code(text, changed=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    if changed:
        highlight_run(run)
    return p

def add_table_row(table, cells, changed_cols=None):
    changed_cols = changed_cols or []
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        if i in changed_cols:
            highlight_run(run)
    return row

# ═══════════════════════════════════════════════════════════════════════════════
# LEGEND
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Article: MCP Servers vs Agent Skills — Updated Version", level=1)
p = doc.add_paragraph()
r1 = p.add_run("Highlight legend:  ")
r1.bold = True
r2 = p.add_run("Yellow = new or changed content  ")
highlight_run(r2)
r2.bold = True
r3 = p.add_run("  Red strikethrough = removed content  ")
strike_run(r3)
r3.font.color.rgb = RGBColor(192, 0, 0)
r3.bold = True
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Introduction", level=1)
add_para([
    ("Versatility in agentic coding is increasing as new tools such as Model Context Protocol (MCP) "
     "servers and Agent Skills become more common. At the same time, many developers ask the same "
     "question when building AI applications: should they use MCP servers or Agent Skills? The "
     "important thing is understanding what each approach does well and choosing the one that fits "
     "your use case. In this post, we\u2019ll explain what MCP servers and Agent Skills are and how "
     "they differ, including architecture diagrams and technical details. In the later sections, "
     "we\u2019ll also walk through how to use Weaviate Agent Skills with Claude Code to build a "
     "\u201cSemantic Movie Discovery\u201d application with several useful features. Let\u2019s get started!", False)
])

# ═══════════════════════════════════════════════════════════════════════════════
# MCP
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Understanding MCP", level=1)
add_para([("The Model Context Protocol (MCP) is an open standard introduced by Anthropic that enables "
           "Large Language Models (LLMs) to interact with external systems such as data sources, APIs "
           "and services. MCP provides a structured way for an AI agent to connect to compliant tools "
           "through a single interface instead of requiring custom integrations for each service.", False)])

add_para([("Image showing the MCP Architecture", False, False, True)])
add_heading("MCP Architecture", level=2)
add_para([("The MCP system operates on a client\u2013server model and consists of three main components.", False)])
add_para([("Host: ", False, True), ("the application that runs the AI model and provides the environment where the agent operates.", False)])
add_para([("Client: ", False, True), ("the protocol connector inside the host that handles communication between the model and MCP servers.", False)])
add_para([("Server: ", False, True), ("an external service that exposes tools, resources, or prompts that the agent can access.", False)])

add_heading("MCP and Agentic Coding", level=2)
add_para([("Before MCP, each AI tool required custom integrations for every external service it wanted "
           "to connect to. MCP simplifies this process by introducing a shared protocol that multiple "
           "agents and tools can use. Developers can now expose capabilities through an MCP server once "
           "and allow any compatible agent to access them without building separate integrations for each system.", False)])

# ═══════════════════════════════════════════════════════════════════════════════
# AGENT SKILLS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Understanding Agent Skills", level=1)
add_para([("Agent Skills, also introduced by Anthropic, provide developers with a simple way to extend "
           "AI coding agents without running MCP servers. An Agent Skill is a structured configuration "
           "file, usually written in YAML or JSON, that defines capabilities, parameter schemas and "
           "natural-language instructions describing how the agent should use those capabilities. AI "
           "tools such as Claude Code read these files at session start and load the skills directly "
           "into the agent\u2019s working context without requiring an additional runtime.", False)])

add_para([("Image showing the Agent Skills Architecture", False, False, True)])
add_heading("How Agent Skills Work", level=2)
add_para([("When Claude Code detects a skill file in the project directory (typically under "
           ".claude/skills/), it loads the manifest into the agent\u2019s context at the beginning "
           "of the session. The skill definition describes available capabilities, how to invoke them "
           "correctly and when to prefer one approach over another. Because the instructions are written "
           "in natural language alongside parameter schemas, the agent can reason about how to use the skill.", False)])
add_para([("Skills are portable across repositories. If a developer commits a skill YAML file to a "
           "repository, any collaborator who clones the project and opens it in Claude Code "
           "automatically gains access to the same capabilities without additional setup.", False)])
add_para([("MCP and Agent Skills solve different problems in agent systems. MCP provides a standardized "
           "way for AI agents to connect to external tools, APIs, databases and services through a "
           "client\u2013server architecture with structured schemas. Agent Skills extend the agent\u2019s "
           "capabilities through configuration files that define workflows, instructions and parameter "
           "schemas without requiring a running server. In simple terms, MCP enables agents to access "
           "external systems, while Agent Skills define how agents perform tasks or workflows within "
           "their environment.", False)])

# ═══════════════════════════════════════════════════════════════════════════════
# WEAVIATE AGENT SKILLS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Weaviate Agent Skills", level=1)
add_para([("Weaviate has released an official set of Agent Skills designed for use with Claude Code "
           "and other compatible agent-based development environments like Cursor, Antigravity, "
           "Windsurf and more. These skills provide structured access to Weaviate vector databases, "
           "allowing agents to perform common operations such as search, querying, schema inspection "
           "data exploration and collection management.", False)])
add_para([("The repository includes ready-to-use skill definitions for tasks like semantic, hybrid "
           "and keyword search, along with natural language querying through the Query Agent. It also "
           "supports workflows such as creating collections, importing data and fetching filtered "
           "results, enabling agents to interact with Weaviate data and perform multi-step retrieval "
           "tasks more effectively.", False)])

add_heading("Weaviate Ecosystem Tools and Features", level=2)
add_heading("Agent Skills and Vector Databases", level=2)
add_para([("AI coding agents face difficulties when working with vector databases. Vector database APIs "
           "provide extensive capabilities, including basic \u201ckey\u2013value\u201d retrieval, "
           "single-vector near-text searches, multimodal near-image searches, hybrid BM25-plus-vector "
           "search, generative modules and multi-tenant system support. Without structured guidance, "
           "even a capable coding agent may produce suboptimal queries: correct syntax but the wrong "
           "search strategy, missing parameters or failure to use powerful features like the Weaviate "
           "Query Agent. Weaviate Agent Skills address this by providing correct usage patterns, "
           "parameter recommendations and decision logic, enabling coding agents to generate "
           "production-ready code from their initial attempts.", False)])
add_para([("The Weaviate Agent Skills repository is organized into two main parts:", False)])
add_para([("Image: Overview of Weaviate Agent Skills", False, False, True)])
add_para([("Weaviate Skill (skills/weaviate): ", False, True),
          ("Focused scripts for tasks such as schema inspection, data ingestion and vector search. "
           "Agents use these while writing application logic or backend code.", False)])
add_para([("Cookbooks Skill (skills/weaviate-cookbooks): ", False, True),
          ("End-to-end project examples that combine tools such as FastAPI, Next.js and Weaviate to "
           "demonstrate full application workflows.", False)])
add_para([("Weaviate Agent Skills work with several development environments, including Claude Code, "
           "Cursor, GitHub Copilot, VS Code and Gemini CLI. When connected to a Weaviate Cloud "
           "instance, agents can directly interact with database modules and perform search, data "
           "management and retrieval tasks.", False)])
add_para([("To evaluate how effective Weaviate Agent Skills really are, let\u2019s build a small "
           "project and see how they accelerate RAG and agentic application development with Claude Code.", False)])

# ═══════════════════════════════════════════════════════════════════════════════
# BUILDING THE APP
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Building the Semantic Movie Discovery Application", level=1)
add_para([("We will build a Movie Discovery App that takes a natural-language description and returns "
           "the three most semantically similar movies from a Weaviate collection. In the process, we "
           "will explore Weaviate capabilities such as multimodal storage, named vector search, "
           "generative AI (RAG) and the Query Agent in action with Claude Code, showing how these "
           "Agentic tools help you build applications faster.", False)])

add_heading("Prerequisites", level=2)
add_para([("Python 3.10 or higher", False)])
add_para([("Weaviate Cloud \u2013 Create a free cluster and obtain an API key.", False)])
add_para([("OpenAI API key \u2013 Required for RAG features.", False)])
add_para([("TMDB API key \u2013 Used to fetch movie metadata and poster images ", False),
          ("during the one-time ingestion step only", True, True),
          ("; not required to run the live app.", True)])
add_para([("Access to Claude Code", False)])
add_para([("Node.js 18+ and npm \u2013 Required to run the Next.js frontend.", True, False, False, False, None)])

# STEP 1
add_heading("Step 1: Project Setup", level=2)
add_para([("Create a movie-discovery-app folder", False)])
add_code("mkdir movie-discovery-app")
add_para([("Create and activate a Python virtual environment in the folder", False)])
add_code("cd movie-discovery-app\npy -m venv venv && source venv\\Scripts\\activate.bat")
add_para([("Install Python dependencies", True)])
add_code("pip install weaviate-client==4.20.1 fastapi uvicorn[standard] openai weaviate-agents>=1.3.0 requests python-dotenv", changed=True)
add_para([("Install Node.js dependencies for the frontend", True)])
add_code("cd frontend && npm install", changed=True)
add_para([("Now create a .env file at the project root. Add the following parameters:", False)])
add_code("WEAVIATE_URL=your-cluster-host-without-https\nWEAVIATE_API_KEY=your-api-key\nOPENAI_API_KEY=your-openai-key\nTMDB_API_KEY=your-tmdb-api-key   # needed for ingestion only")

add_para([("After signing up for Weaviate, click the Create Cluster button to start a new cluster. "
           "Click \u201cHow to Connect\u201d to view the required Weaviate connection parameters.", False)])
add_para([("Now that everything is set up, we can connect Weaviate Cloud with Claude Code by running "
           "claude in your project terminal.", False)])
add_code('4. Write and run `check_modules.py` that connects using `weaviate.connect_to_weaviate_cloud` with `skip_init_checks=True`, loads credentials from `.env` with `python-dotenv`, and prints the full JSON list of enabled Weaviate modules. Run it with `venv/Scripts/python check_modules.py`.')

# STEP 2
add_heading("Step 2: Create a Weaviate Movie Collection", level=2)
add_para([("A Weaviate collection schema named \u201cMovie\u201d needs to be created. To avoid errors, "
           "the schema should follow these requirements:", False)])
add_para([("Use the Weaviate REST API directly with the requests library.", False)])
add_para([("When using multi2multivec-weaviate, the vector format vectorConfig must be used.", False)])
add_para([("Use text2vec-weaviate for text fields, as this is architecturally cleaner.", False)])
add_para([("Use the following prompt to generate a Weaviate collection schema in a create_schema.py file:", False)])
add_code('Write and run a Python script `create_schema.py` that creates a Weaviate collection named `Movie`.\nFollow these constraints exactly...\n[full prompt as in original article]')

# STEP 3
add_heading("Step 3: Ingest Sample Movie Data", level=2)
add_para([("In this step, we encode the required movie data and upload it to Weaviate Cloud. The "
           "sample dataset contains ", False),
          ("100 movies", False, True),
          (". The TMDB API provides structured endpoints for retrieving the required metadata, "
           "while poster images are downloaded from the TMDB public CDN and stored as base64 blobs. "
           "This is where Weaviate Agent Skills come into action, helping ingest data from the movie "
           "API into the database.", False)])
add_para([("Use this prompt to ingest movie data:", False)])
add_code('Write and run a Python script `ingest_movies.py` that fetches 100 movies from the TMDB API\nand ingests them into the `Movie` Weaviate collection.\n[full prompt as in original article]')
add_para([("Note: Weaviate multimodal embedding ensures that multi2multivec-weaviate handles image "
           "embedding, while text2vec-weaviate handles text embedding.", False, False, True)])

# STEP 4 — CHANGED
add_heading("Step 4: Building the FastAPI Backend and Next.js Frontend with a Weaviate Query Agent Chat Layer", level=2, changed=True)
add_para([("The app uses a ", False),
          ("two-layer architecture", True, True),
          (": a ", True),
          ("FastAPI backend", True, True),
          (" that exposes REST endpoints and a ", True),
          ("Next.js (TypeScript) frontend", True, True),
          (" that renders the UI. The backend connects directly to Weaviate Cloud and the Weaviate "
           "Query Agent. The frontend communicates with the backend over HTTP.", True)])

add_para([("The app is organized into two views accessed via a collapsible sidebar:", True)])
add_para([("Search view: ", True, True),
          ("performs semantic search and RAG using Weaviate named vectors.", True)])
add_para([("Chat view: ", True, True),
          ("handles multi-turn conversations through the Weaviate Query Agent.", True)])

# Architecture table
add_para([("The table below describes the architecture of the built app:", True)])
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Layer"
hdr[1].text = "Component"
hdr[2].text = "Role"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
        highlight_run(run)

rows = [
    ("Backend", "backend.py (FastAPI)", "REST API on port 8000; routes: GET /health, GET /search, POST /ai/explain, POST /ai/plan, POST /chat"),
    ("Frontend", "Next.js + TypeScript (port 3000)", "Single-page app with sidebar navigation"),
    ("", "SearchView.tsx", "Semantic search (near_text), AI explanations (single_prompt), Movie Night Planner (grouped_task)"),
    ("", "MovieCard.tsx", "Renders base64 poster inline, watchlist add/remove button"),
    ("", "ChatView.tsx", "Multi-turn Query Agent chat with collapsible source citations"),
    ("", "AppSidebar.tsx", "Navigation (Search/Chat), Weaviate logo + feature summary, watchlist manager with .txt export"),
]
for cells in rows:
    add_table_row(tbl, cells, changed_cols=[0, 1, 2])

doc.add_paragraph()

# Sidebar feature detail
add_para([("The sidebar includes:", True)])
add_para([("Weaviate logo and a one-line description of the app", True)])
add_para([("Feature list: text2vec-weaviate, multi2multivec-weaviate, generative-openai, near_text "
           "search, single_prompt RAG, grouped_task RAG, Query Agent", True)])
add_para([("Watchlist manager: add/remove movies, export to .txt", True)])

# Prompt block
add_para([("Use the following prompts with Claude Code to generate the backend and frontend:", True)])
add_para([("Prompt 1 \u2014 Backend:", True, True)])
add_code(
    'Create `backend.py`: a FastAPI app with CORS enabled for localhost:3000.\n'
    'Connect to Weaviate Cloud using credentials from .env with skip_init_checks=True.\n'
    'Implement these routes:\n'
    '  GET  /health                  \u2192 {"status": "ok"}\n'
    '  GET  /search?q=...&limit=3    \u2192 near_text on text_vector, return title/description/release_year/poster\n'
    '  POST /ai/explain              \u2192 generate.near_text with single_prompt\n'
    '  POST /ai/plan                 \u2192 generate.near_text with grouped_task\n'
    '  POST /chat                    \u2192 QueryAgent.ask() with full message history\n'
    'Run with: uvicorn backend:app --reload --port 8000',
    changed=True
)
add_para([("Prompt 2 \u2014 Frontend:", True, True)])
add_code(
    'Create a Next.js TypeScript app in the frontend/ folder.\n'
    'Components needed:\n'
    '  page.tsx        \u2014 SidebarProvider layout, view state (search | chat)\n'
    '  SearchView.tsx  \u2014 search input, MovieCard grid, AI explain and plan buttons\n'
    '  MovieCard.tsx   \u2014 poster image, title, year, description, watchlist button\n'
    '  ChatView.tsx    \u2014 message bubbles, source citations, clear chat\n'
    '  AppSidebar.tsx  \u2014 navigation, Weaviate logo + feature list, watchlist + export\n'
    'Backend base URL from NEXT_PUBLIC_BACKEND_HOST env var (default localhost:8000)\n'
    'Run with: npm run dev',
    changed=True
)

# Starting the app — CHANGED
add_heading("Starting the Application", level=2, changed=True)
add_para([("Start both processes in separate terminals:", True)])
add_code("# Terminal 1 \u2014 Backend\nuvicorn backend:app --reload --port 8000", changed=True)
add_code("# Terminal 2 \u2014 Frontend\ncd frontend && npm run dev", changed=True)
add_para([("The app is available at ", True),
          ("http://localhost:3000", True, True),
          (". The API runs at ", True),
          ("http://localhost:8000", True, True),
          (".", True)])
add_para([("After this, Claude Code will automatically build the app by adding relevant files and "
           "start both servers. You can start using the application immediately.", False)])

# ═══════════════════════════════════════════════════════════════════════════════
# DEMO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Congratulations!", level=1)
add_para([("You\u2019ve completed the project without needing to do much manual configuration or coding.", False)])

add_heading("Demo", level=2)
add_para([("So far we have used Weaviate Agent Skills with Claude Code to build a Semantic Movie "
           "Discovery Application with an OpenAI API key and TMDB API.", False)])
add_para([("Screen Recording 2026-03-05 193106.mp4", False, False, True)])

add_para([("The Movie Discovery app we built includes the following features:", False)])
add_para([("Semantic search: ", False, True),
          ("Describe a mood or theme and retrieve matching movies using vector-based search (near_text).", False)])
add_para([("AI explanations: ", False, True),
          ("Generate per-movie summaries using RAG with single_prompt.", False)])
add_para([("Movie Night Planner: ", False, True),
          ("Create a viewing order, snack pairings and a theme summary using grouped_task.", False)])
add_para([("Conversational chat: ", False, True),
          ("Ask questions about the movie collection through a chat interface powered by the "
           "Weaviate Query Agent, with source citations.", False)])
add_para([("Watchlist: ", False, True),
          ("Save movies during your session and export the list as a .txt file.", False)])
add_para([("Poster download: Download any movie poster as a JPEG.", False, False, False, True)])  # strikethrough = removed

# ═══════════════════════════════════════════════════════════════════════════════
# WHAT'S NEXT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("What\u2019s Next?", level=1)
add_para([("You could add genre filters to target specific audiences and better meet your movie choice. "
           "You could also include a hybrid search feature that incorporates keyword-heavy queries and "
           "image search. You can take your app even further by getting up to speed with Weaviate\u2019s "
           "latest releases and becoming familiar with features such as server-side batching, async "
           "replication improvements, Object TTL and many more. To explore further, check out the "
           "latest Weaviate releases and join the discussion on the community forum.", False)])

# ═══════════════════════════════════════════════════════════════════════════════
# WEAVIATE IN ACTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Weaviate Agent Skills in Action", level=1)
add_para([("In this project, we saw how Weaviate modules were used in the application:", False)])
add_para([("Text2vec-weaviate: ", False, True), ("Responsible for text embeddings.", False)])
add_para([("Multi2multivec-weaviate: ", False, True), ("Responsible for embedding images.", False)])
add_para([("Generative-openai: ", False, True), ("Integrates GPT directly into the query workflow.", False)])
add_para([("Weaviate Query Agent: ", False, True),
          ("A higher-level abstraction that accepts natural language queries, decides the best query "
           "method, executes queries, synthesizes results and returns answers.", False)])
add_para([("Weaviate Agent Skills help in shipping faster and more accurate RAG applications. Backend "
           "development tasks such as schema inspection, data ingestion and search operations are "
           "automated and optimized. Ultimately, this helps developers save valuable development time.", False)])

# ═══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Conclusion", level=1)
add_para([("Both MCP servers and Agent Skills provide useful patterns for building AI-powered "
           "applications. MCP servers are well suited for exposing external tools and services through "
           "a standardized interface, while Agent Skills focus on guiding coding agents with structured "
           "workflows and best practices.", False)])
add_para([("In this tutorial, we demonstrated how Weaviate Agent Skills can simplify development by "
           "helping Claude Code generate correct database queries, ingestion pipelines and search logic. "
           "By combining vector search, multimodal storage and generative capabilities, we built a "
           "semantic movie discovery application with minimal manual setup.", False)])
add_para([("As agentic development environments continue to evolve, tools like MCP servers and Agent "
           "Skills will likely be used together. The key is understanding where each approach fits and "
           "selecting the one that best supports your application architecture. Happy building.", False)])

# ═══════════════════════════════════════════════════════════════════════════════
# RESOURCES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Resources", level=1)
add_para([("Model Context Protocol", False)])
add_para([("Weaviate Agent Skills", False)])
add_para([("Claude Code", False)])
add_para([("GitHub Repository for the Movie Discovery App", False)])

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("article_updated.docx")
print("Saved: article_updated.docx")
